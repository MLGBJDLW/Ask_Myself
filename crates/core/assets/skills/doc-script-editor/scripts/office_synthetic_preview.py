"""Clearly labelled structural previews that never count as final Office render evidence."""

from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _canvas(title: str):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (1280, 720), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1280, 48), fill="#18233A")
    draw.text((22, 15), f"SYNTHETIC STRUCTURAL PREVIEW — {title}", fill="white")
    return image, draw


def _pptx(path: Path, output: Path) -> list[dict[str, Any]]:
    scripts = Path(__file__).resolve().parents[2] / "pptx-presentation-design" / "scripts"
    if str(scripts) not in sys.path:
        sys.path.insert(0, str(scripts))
    from pptx_audit import audit

    profile = audit(path)
    cx = max(int(profile.get("slide_size", {}).get("cx", 1)), 1)
    cy = max(int(profile.get("slide_size", {}).get("cy", 1)), 1)
    files = []
    for slide in profile.get("slide_details", []):
        image, draw = _canvas(f"slide {slide['index']}")
        for shape in slide.get("shape_details", []):
            bounds = shape.get("bounds", {})
            if any(bounds.get(key) is None for key in ("x", "y", "cx", "cy")):
                continue
            left = 30 + int(1220 * int(bounds["x"]) / cx)
            top = 70 + int(620 * int(bounds["y"]) / cy)
            right = left + max(2, int(1220 * int(bounds["cx"]) / cx))
            bottom = top + max(2, int(620 * int(bounds["cy"]) / cy))
            draw.rectangle((left, top, right, bottom), outline="#2457D6", width=2)
            label = (shape.get("text") or shape.get("shapeName") or shape.get("kind") or "shape")[:80]
            draw.text((left + 4, top + 4), label, fill="#17213A")
        target = output / f"slide-{int(slide['index']):03d}.png"
        image.save(target)
        files.append({"surface": f"slide:{slide['slide_id']}", "path": str(target), "sha256": _sha256(target)})
    return files


def _xlsx(path: Path, output: Path) -> list[dict[str, Any]]:
    import openpyxl

    workbook = openpyxl.load_workbook(path, data_only=False, read_only=True)
    files = []
    try:
        for index, worksheet in enumerate(
            (sheet for sheet in workbook.worksheets if sheet.sheet_state == "visible"), start=1
        ):
            image, draw = _canvas(f"worksheet {worksheet.title}")
            left, top, cell_width, cell_height = 35, 78, 118, 28
            for row in range(1, min(worksheet.max_row, 20) + 1):
                for column in range(1, min(worksheet.max_column, 10) + 1):
                    x0 = left + (column - 1) * cell_width
                    y0 = top + (row - 1) * cell_height
                    draw.rectangle((x0, y0, x0 + cell_width, y0 + cell_height), outline="#CBD4E4")
                    value = worksheet.cell(row, column).value
                    if value is not None:
                        draw.text((x0 + 4, y0 + 7), str(value)[:18], fill="#17213A")
            target = output / f"sheet-{index:03d}.png"
            image.save(target)
            files.append({"surface": f"worksheet:{worksheet.title}", "path": str(target), "sha256": _sha256(target)})
    finally:
        workbook.close()
    return files


def _docx(path: Path, output: Path) -> list[dict[str, Any]]:
    import docx

    document = docx.Document(path)
    image, draw = _canvas("Word content flow")
    y = 72
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            y += 10
            continue
        draw.text((60, y), f"[{paragraph.style.name}] {text}"[:150], fill="#17213A")
        y += 26
        if y > 620:
            break
    for table in document.tables[:4]:
        if y > 620:
            break
        draw.rectangle((55, y, 1225, min(680, y + 32 * max(1, len(table.rows)))), outline="#2457D6", width=2)
        draw.text((65, y + 8), f"Table {len(table.rows)}x{len(table.columns)}", fill="#17213A")
        y += 32 * max(1, len(table.rows)) + 12
    target = output / "document-flow-001.png"
    image.save(target)
    return [{"surface": "document-flow", "path": str(target), "sha256": _sha256(target)}]


def create_synthetic_preview(path: Path, output: Path) -> dict[str, Any]:
    output.mkdir(parents=True, exist_ok=False)
    suffix = path.suffix.lower()
    if suffix in {".pptx", ".pptm", ".potx", ".potm"}:
        files = _pptx(path, output)
    elif suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        files = _xlsx(path, output)
    elif suffix in {".docx", ".docm", ".dotx", ".dotm"}:
        files = _docx(path, output)
    else:
        raise ValueError(f"unsupported Office preview format: {suffix}")
    manifest = {
        "kind": "officeSyntheticPreview",
        "artifactSha256": _sha256(path),
        "isFinalRenderEvidence": False,
        "warning": "Approximate structural preview; never substitutes for final Office/LibreOffice rendering.",
        "files": files,
    }
    manifest_path = output / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    manifest["manifestPath"] = str(manifest_path)
    return manifest

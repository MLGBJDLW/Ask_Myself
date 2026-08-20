#!/usr/bin/env python3
"""Deterministic DOCX Spec v2 renderer with layout and accessibility evidence."""

from __future__ import annotations

import argparse
import json
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any


PRESETS: dict[str, dict[str, Any]] = {
    "executive": {
        "bodyFont": "Aptos",
        "headingFont": "Aptos Display",
        "bodySize": 10.5,
        "titleSize": 28,
        "primary": "17365D",
        "accent": "2F75B5",
        "tableHeader": "D9EAF7",
        "callout": "EAF2F8",
        "paragraphAfter": 6,
    },
    "technical": {
        "bodyFont": "Aptos",
        "headingFont": "Aptos Display",
        "bodySize": 10,
        "titleSize": 26,
        "primary": "263238",
        "accent": "00796B",
        "tableHeader": "D7EEEA",
        "callout": "E8F5F2",
        "paragraphAfter": 5,
    },
    "proposal": {
        "bodyFont": "Aptos",
        "headingFont": "Georgia",
        "bodySize": 11,
        "titleSize": 30,
        "primary": "3B2F2F",
        "accent": "A65A3A",
        "tableHeader": "F2E3DA",
        "callout": "FAF1EC",
        "paragraphAfter": 7,
    },
    "memo": {
        "bodyFont": "Aptos",
        "headingFont": "Aptos",
        "bodySize": 10.5,
        "titleSize": 22,
        "primary": "1F2937",
        "accent": "4B5563",
        "tableHeader": "E5E7EB",
        "callout": "F3F4F6",
        "paragraphAfter": 5,
    },
}

ROOT_KEYS = {
    "schemaVersion", "preset", "title", "subtitle", "author", "language", "template",
    "clearTemplateBody", "page", "tokens", "header", "footer", "blocks",
}
BLOCK_KEYS = {
    "heading": {"type", "text", "level"},
    "paragraph": {"type", "text", "style", "keepWithNext"},
    "bullets": {"type", "items"},
    "numbered": {"type", "items"},
    "table": {"type", "headers", "rows", "columnWidths", "repeatHeader", "allowRowBreaks", "caption"},
    "image": {"type", "path", "width", "altText", "caption"},
    "callout": {"type", "text", "kind"},
    "pageBreak": {"type"},
    "sectionBreak": {"type", "orientation", "page"},
}
PAGE_KEYS = {"orientation", "marginTop", "marginBottom", "marginLeft", "marginRight"}
HEADER_FOOTER_KEYS = {"text", "pageNumber", "differentFirstPage"}
HEX_RE = re.compile(r"^[0-9A-Fa-f]{6}$")


class DocxSpecError(ValueError):
    pass


def _validate_keys(value: dict[str, Any], allowed: set[str], location: str) -> None:
    unknown = sorted(set(value) - allowed)
    if unknown:
        raise DocxSpecError(f"unknown field(s) at {location}: {', '.join(unknown)}")


def validate_spec(spec: dict[str, Any]) -> None:
    _validate_keys(spec, ROOT_KEYS, "$.")
    if spec.get("schemaVersion") != 2:
        raise DocxSpecError("schemaVersion must be 2")
    if str(spec.get("preset", "executive")) not in PRESETS:
        raise DocxSpecError("preset must be executive, technical, proposal, or memo")
    if not isinstance(spec.get("blocks"), list):
        raise DocxSpecError("blocks must be an array")
    if isinstance(spec.get("page"), dict):
        _validate_keys(spec["page"], PAGE_KEYS, "$.page")
    for name in ("header", "footer"):
        if isinstance(spec.get(name), dict):
            _validate_keys(spec[name], HEADER_FOOTER_KEYS, f"$.{name}")
    for index, block in enumerate(spec["blocks"]):
        if not isinstance(block, dict):
            raise DocxSpecError(f"block {index} must be an object")
        kind = str(block.get("type", ""))
        if kind not in BLOCK_KEYS:
            raise DocxSpecError(f"unsupported block type at $.blocks[{index}]: {kind or '<missing>'}")
        _validate_keys(block, BLOCK_KEYS[kind], f"$.blocks[{index}]")
        if kind == "heading" and not 1 <= int(block.get("level", 1)) <= 6:
            raise DocxSpecError(f"heading level out of range at $.blocks[{index}]")
        if kind in {"bullets", "numbered"} and not isinstance(block.get("items"), list):
            raise DocxSpecError(f"{kind} items must be an array at $.blocks[{index}]")
        if kind == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if not isinstance(headers, list) or not isinstance(rows, list) or not all(isinstance(row, list) for row in rows):
                raise DocxSpecError(f"table headers/rows are invalid at $.blocks[{index}]")
            columns = len(headers) or max((len(row) for row in rows), default=0)
            if columns == 0 or any(len(row) != columns for row in rows):
                raise DocxSpecError(f"table rows must have a stable non-zero column count at $.blocks[{index}]")
            widths = block.get("columnWidths")
            if widths is not None and (not isinstance(widths, list) or len(widths) != columns):
                raise DocxSpecError(f"columnWidths must match table columns at $.blocks[{index}]")
        if kind == "image" and not str(block.get("altText", "")).strip():
            raise DocxSpecError(f"image altText is required at $.blocks[{index}]")


def _workspace_path(raw: str, workspace_root: Path, *, must_exist: bool = True) -> Path:
    path = Path(raw).expanduser()
    resolved = (path if path.is_absolute() else workspace_root / path).resolve()
    try:
        resolved.relative_to(workspace_root)
    except ValueError as error:
        raise DocxSpecError(f"path escapes workspace: {resolved}") from error
    if must_exist and not resolved.exists():
        raise DocxSpecError(f"file not found: {resolved}")
    return resolved


def _hex(value: Any, default: str) -> str:
    text = str(value or default).lstrip("#")
    return text.upper() if HEX_RE.fullmatch(text) else default


def _apply_language(element: Any, language: str) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    rpr = element.get_or_add_rPr()
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    for attribute in ("w:val", "w:eastAsia", "w:bidi"):
        lang.set(qn(attribute), language)


def _configure_styles(document: Any, tokens: dict[str, Any], language: str) -> None:
    from docx.enum.style import WD_STYLE_TYPE  # type: ignore
    from docx.shared import Pt, RGBColor  # type: ignore

    normal = document.styles["Normal"]
    normal.font.name = str(tokens["bodyFont"])
    normal.font.size = Pt(float(tokens["bodySize"]))
    normal.paragraph_format.space_after = Pt(float(tokens["paragraphAfter"]))
    normal.paragraph_format.line_spacing = 1.15
    _apply_language(normal.element, language)
    for level in range(1, 7):
        style = document.styles[f"Heading {level}"]
        style.font.name = str(tokens["headingFont"])
        style.font.color.rgb = RGBColor.from_string(str(tokens["primary"]))
        style.font.size = Pt(max(11, float(tokens["titleSize"]) - level * 3.2))
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        style.paragraph_format.space_after = Pt(4)
        _apply_language(style.element, language)
    if "Nexa Caption" not in document.styles:
        caption = document.styles.add_style("Nexa Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = document.styles["Nexa Caption"]
    caption.font.name = str(tokens["bodyFont"])
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string("666666")


def _apply_page(section: Any, page: dict[str, Any]) -> None:
    from docx.enum.section import WD_ORIENT  # type: ignore
    from docx.shared import Inches  # type: ignore

    orientation = str(page.get("orientation", "portrait"))
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        if section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
    for key, attribute, default in (
        ("marginTop", "top_margin", 0.75),
        ("marginBottom", "bottom_margin", 0.75),
        ("marginLeft", "left_margin", 0.85),
        ("marginRight", "right_margin", 0.85),
    ):
        setattr(section, attribute, Inches(float(page.get(key, default))))


def _clear_body(document: Any) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.rsplit("}", 1)[-1] != "sectPr":
            body.remove(child)


def _add_hyperlink(paragraph: Any, label: str, url: str) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline(paragraph: Any, text: str) -> None:
    link_pattern = re.compile(r"\[([^\]]+)\]\((https?://[^)]+|mailto:[^)]+)\)")
    cursor = 0
    for match in link_pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        _add_hyperlink(paragraph, match.group(1), match.group(2))
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _set_repeat_header(row: Any) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)


def _set_row_split(row: Any, allow: bool) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if not allow and cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    elif allow and cant_split is not None:
        tr_pr.remove(cant_split)


def _shade(cell: Any, color: str) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_fixed_table_layout(table: Any) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    table.autofit = False
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")


def _add_page_field(paragraph: Any) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(f"{{http://www.w3.org/XML/1998/namespace}}space", "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend([begin, instruction, separate, end])
    paragraph._p.append(run)


def _configure_story(story: Any, spec: dict[str, Any] | None) -> None:
    if not spec:
        return
    paragraph = story.paragraphs[0]
    paragraph.clear()
    text = str(spec.get("text", ""))
    if text:
        paragraph.add_run(text)
    if spec.get("pageNumber"):
        if text:
            paragraph.add_run(" · ")
        _add_page_field(paragraph)


def render_docx(spec: dict[str, Any], output: Path, workspace_root: Path) -> dict[str, Any]:
    try:
        import docx  # type: ignore
        from docx.enum.section import WD_SECTION  # type: ignore
        from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Inches, Pt, RGBColor  # type: ignore
    except ImportError as error:
        raise DocxSpecError(f"python-docx is required: {error}") from error

    validate_spec(spec)
    preset = PRESETS[str(spec.get("preset", "executive"))]
    tokens = deepcopy(preset)
    tokens.update(spec.get("tokens") if isinstance(spec.get("tokens"), dict) else {})
    for color_key in ("primary", "accent", "tableHeader", "callout"):
        tokens[color_key] = _hex(tokens.get(color_key), preset[color_key])
    language = str(spec.get("language", "zh-CN"))
    template = spec.get("template")
    document = docx.Document(str(_workspace_path(str(template), workspace_root))) if template else docx.Document()
    if template and spec.get("clearTemplateBody", True):
        _clear_body(document)
    _configure_styles(document, tokens, language)
    _apply_page(document.sections[0], spec.get("page") if isinstance(spec.get("page"), dict) else {})
    document.core_properties.title = str(spec.get("title", ""))
    document.core_properties.author = str(spec.get("author", "Nexa"))

    first_section = document.sections[0]
    _configure_story(first_section.header, spec.get("header") if isinstance(spec.get("header"), dict) else None)
    _configure_story(first_section.footer, spec.get("footer") if isinstance(spec.get("footer"), dict) else None)

    if spec.get("title"):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(spec["title"]))
        run.font.name = str(tokens["headingFont"])
        run.font.size = Pt(float(tokens["titleSize"]))
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(str(tokens["primary"]))
    if spec.get("subtitle"):
        paragraph = document.add_paragraph(str(spec["subtitle"]))
        paragraph.style = document.styles["Subtitle"]

    metrics = {"headings": 0, "paragraphs": 0, "tables": 0, "images": 0, "callouts": 0, "sections": 1}
    for block in spec["blocks"]:
        kind = str(block["type"])
        if kind == "heading":
            paragraph = document.add_heading(level=int(block.get("level", 1)))
            _add_inline(paragraph, str(block.get("text", "")))
            metrics["headings"] += 1
        elif kind == "paragraph":
            paragraph = document.add_paragraph(style=block.get("style") or None)
            _add_inline(paragraph, str(block.get("text", "")))
            paragraph.paragraph_format.keep_with_next = bool(block.get("keepWithNext", False))
            metrics["paragraphs"] += 1
        elif kind in {"bullets", "numbered"}:
            style = "List Bullet" if kind == "bullets" else "List Number"
            for item in block.get("items", []):
                paragraph = document.add_paragraph(style=style)
                _add_inline(paragraph, str(item))
                metrics["paragraphs"] += 1
        elif kind == "table":
            headers = list(block.get("headers", []))
            rows = list(block.get("rows", []))
            columns = len(headers) or len(rows[0])
            table = document.add_table(rows=len(rows) + (1 if headers else 0), cols=columns)
            table.style = "Table Grid"
            _set_fixed_table_layout(table)
            source_rows = ([headers] if headers else []) + rows
            widths = [float(value) for value in block.get("columnWidths", [])]
            if not widths:
                available = (document.sections[-1].page_width - document.sections[-1].left_margin - document.sections[-1].right_margin) / 914400
                widths = [available / columns] * columns
            for row_index, values in enumerate(source_rows):
                row = table.rows[row_index]
                if row_index == 0 and headers and block.get("repeatHeader", True):
                    _set_repeat_header(row)
                _set_row_split(row, bool(block.get("allowRowBreaks", True)))
                for column_index, value in enumerate(values):
                    cell = row.cells[column_index]
                    cell.width = Inches(widths[column_index])
                    cell.text = ""
                    _add_inline(cell.paragraphs[0], str(value))
                    if row_index == 0 and headers:
                        _shade(cell, str(tokens["tableHeader"]))
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
            if block.get("caption"):
                document.add_paragraph(str(block["caption"]), style="Nexa Caption")
            metrics["tables"] += 1
        elif kind == "image":
            image = _workspace_path(str(block.get("path", "")), workspace_root)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shape = paragraph.add_run().add_picture(str(image), width=Inches(float(block.get("width", 5.5))))
            shape._inline.docPr.set("descr", str(block["altText"]))
            shape._inline.docPr.set("title", str(block.get("caption") or block["altText"]))
            if block.get("caption"):
                caption = document.add_paragraph(str(block["caption"]), style="Nexa Caption")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            metrics["images"] += 1
        elif kind == "callout":
            table = document.add_table(rows=1, cols=1)
            _set_fixed_table_layout(table)
            cell = table.cell(0, 0)
            palette = {"info": tokens["callout"], "success": "E2F0D9", "warning": "FFF2CC", "risk": "FCE4D6"}
            _shade(cell, str(palette.get(str(block.get("kind", "info")), tokens["callout"])))
            cell.text = ""
            _add_inline(cell.paragraphs[0], str(block.get("text", "")))
            metrics["callouts"] += 1
        elif kind == "pageBreak":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            section = document.add_section(WD_SECTION.NEW_PAGE)
            page = block.get("page") if isinstance(block.get("page"), dict) else {}
            page = {**page, "orientation": block.get("orientation", page.get("orientation", "portrait"))}
            _apply_page(section, page)
            _configure_story(section.header, spec.get("header") if isinstance(spec.get("header"), dict) else None)
            _configure_story(section.footer, spec.get("footer") if isinstance(spec.get("footer"), dict) else None)
            metrics["sections"] += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "kind": "docxSpecRender",
        "schemaVersion": 2,
        "path": str(output),
        "preset": str(spec.get("preset", "executive")),
        "language": language,
        "metrics": metrics,
        "qualityEvidence": {
            "deterministicStyles": True,
            "fixedTableGeometry": True,
            "imageAltTextRequired": True,
            "renderedPagesInspected": False,
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Render DOCX Spec v2")
    parser.add_argument("--spec", required=True)
    parser.add_argument("--path", required=True)
    parser.add_argument("--workspace-root", default=None)
    args = parser.parse_args()
    root = Path(args.workspace_root).resolve() if args.workspace_root else Path.cwd().resolve()
    try:
        spec = json.loads(Path(args.spec).read_text(encoding="utf-8"))
        if not isinstance(spec, dict):
            raise DocxSpecError("spec root must be an object")
        result = render_docx(spec, _workspace_path(args.path, root, must_exist=False), root)
    except (OSError, json.JSONDecodeError, DocxSpecError) as error:
        print(f"DOCX_RENDER_FAILED: {error}", file=sys.stderr)
        return 1
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

#!/usr/bin/env python3
"""Tests for basic DOCX review lifecycle."""

from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import docx_review_editor
import docx_audit


class DocxReviewEditorTests(unittest.TestCase):
    def _source(self, root: Path, text: str = "Review target here") -> Path:
        import docx

        path = root / "source.docx"
        document = docx.Document()
        document.add_paragraph(text)
        document.save(path)
        return path

    def test_comment_add_extract_and_strip_lifecycle(self) -> None:
        import docx

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = self._source(root)
            commented = root / "commented.docx"
            stripped = root / "stripped.docx"

            result = docx_review_editor.patch_docx_reviews(source, commented, [{
                "op": "add_comment",
                "find": "target",
                "comment": "Confirm this wording.",
                "author": "Reviewer",
                "initials": "RV",
            }])
            comments = docx_review_editor.extract_comments(commented)

            self.assertIn("word/comments.xml", result["changedParts"])
            self.assertEqual("Confirm this wording.", comments["comments"][0]["text"])
            self.assertEqual("Reviewer", comments["comments"][0]["author"])
            self.assertEqual("Review target here", docx.Document(commented).paragraphs[0].text)
            with zipfile.ZipFile(commented) as archive:
                document_xml = archive.read("word/document.xml")
                self.assertIn(b"commentRangeStart", document_xml)
                self.assertIn(b"commentReference", document_xml)

            docx_review_editor.patch_docx_reviews(commented, stripped, [{"op": "strip_comments"}])
            self.assertEqual([], docx_review_editor.extract_comments(stripped)["comments"])
            self.assertEqual("Review target here", docx.Document(stripped).paragraphs[0].text)
            with zipfile.ZipFile(stripped) as archive:
                self.assertNotIn("word/comments.xml", archive.namelist())

    def test_comment_replies_and_thread_resolution_use_native_metadata_parts(self) -> None:
        import docx

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = self._source(root)
            threaded = root / "threaded.docx"

            result = docx_review_editor.patch_docx_reviews(source, threaded, [
                {
                    "op": "add_comment",
                    "find": "target",
                    "comment": "Parent review",
                    "author": "Reviewer A",
                },
                {
                    "op": "reply_comment",
                    "commentId": "0",
                    "comment": "Reply review",
                    "author": "Reviewer B",
                },
                {"op": "resolve_comment", "commentId": "0", "resolved": True},
            ])

            self.assertEqual(["0", "1"], result["operations"][2]["detail"]["affectedCommentIds"])
            comments = docx_review_editor.extract_comments(threaded)["comments"]
            self.assertEqual(2, len(comments))
            self.assertIsNone(comments[0]["parentId"])
            self.assertEqual("0", comments[1]["parentId"])
            self.assertTrue(all(comment["resolved"] for comment in comments))
            self.assertTrue(all(comment["paraId"] for comment in comments))
            self.assertTrue(all(comment["durableId"] for comment in comments))
            self.assertEqual("Review target here", docx.Document(threaded).paragraphs[0].text)
            audit = docx_audit.audit(threaded)
            self.assertEqual(1, audit["comment_replies"])
            self.assertEqual(2, audit["resolved_comments"])
            with zipfile.ZipFile(threaded) as archive:
                for part in (
                    "word/comments.xml",
                    "word/commentsExtended.xml",
                    "word/commentsIds.xml",
                    "word/commentsExtensible.xml",
                ):
                    self.assertIn(part, archive.namelist())
                extended = archive.read("word/commentsExtended.xml")
                self.assertIn(b"paraIdParent", extended)
                self.assertEqual(2, extended.count(b'done="1"'))
                comments_xml = archive.read("word/comments.xml")
                self.assertIn(b'mc:Ignorable="w14"', comments_xml)
                self.assertIn(b'xmlns:w14=', comments_xml)
                self.assertNotIn(b"Ignorable=", archive.read("word/document.xml"))

    def test_comment_add_strip_and_recreate_are_sequential_in_one_request(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = self._source(root)
            stripped = root / "add-then-strip.docx"
            docx_review_editor.patch_docx_reviews(source, stripped, [
                {"op": "add_comment", "find": "target", "comment": "Transient"},
                {"op": "strip_comments"},
            ])
            self.assertEqual([], docx_review_editor.extract_comments(stripped)["comments"])
            with zipfile.ZipFile(stripped) as archive:
                self.assertFalse(any(name.startswith("word/comments") for name in archive.namelist()))
                self.assertNotIn(b"commentReference", archive.read("word/document.xml"))

            existing = root / "existing.docx"
            docx_review_editor.patch_docx_reviews(source, existing, [{
                "op": "add_comment",
                "find": "target",
                "comment": "Old review",
            }])
            recreated = root / "strip-then-add.docx"
            docx_review_editor.patch_docx_reviews(existing, recreated, [
                {"op": "strip_comments"},
                {
                    "op": "add_comment",
                    "find": "target",
                    "comment": "New review",
                    "author": "Replacement reviewer",
                },
            ])
            comments = docx_review_editor.extract_comments(recreated)["comments"]
            self.assertEqual(1, len(comments))
            self.assertEqual("New review", comments[0]["text"])
            self.assertEqual("Replacement reviewer", comments[0]["author"])
            self.assertEqual(0, docx_audit.audit(recreated)["comment_replies"])

    def test_accept_reject_fail_closed_for_unhandled_revision_kinds(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = self._source(root, "Moved content")
            complex_redline = root / "complex-redline.docx"
            with zipfile.ZipFile(source) as input_archive, zipfile.ZipFile(complex_redline, "w") as output_archive:
                for info in input_archive.infolist():
                    payload = input_archive.read(info.filename)
                    if info.filename == "word/document.xml":
                        document = ET.fromstring(payload)
                        paragraph = next(item for item in document.iter() if item.tag.endswith("}p"))
                        move = ET.SubElement(paragraph, f"{{{docx_review_editor.W_NS}}}moveFrom")
                        run = ET.SubElement(move, f"{{{docx_review_editor.W_NS}}}r")
                        ET.SubElement(run, f"{{{docx_review_editor.W_NS}}}t").text = "Moved"
                        payload = ET.tostring(document, encoding="utf-8", xml_declaration=True)
                    output_archive.writestr(info, payload)

            with self.assertRaisesRegex(docx_review_editor.DocxReviewError, "unsupported.*moveFrom"):
                docx_review_editor.patch_docx_reviews(
                    complex_redline,
                    root / "unsafe-accepted.docx",
                    [{"op": "accept_changes"}],
                )
            audit = docx_audit.audit(complex_redline)
            self.assertEqual(["moveFrom"], audit["tracked_changes"]["unsupported_for_accept_reject"])

    def test_unknown_markup_compatibility_prefix_fails_closed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = self._source(root)
            future = root / "future-word.docx"
            with zipfile.ZipFile(source) as input_archive, zipfile.ZipFile(future, "w") as output_archive:
                for info in input_archive.infolist():
                    payload = input_archive.read(info.filename)
                    if info.filename == "word/document.xml":
                        document = ET.fromstring(payload)
                        document.set(
                            f"{{{docx_review_editor.MC_NS}}}Ignorable",
                            "w14 w99",
                        )
                        payload = ET.tostring(document, encoding="utf-8", xml_declaration=True)
                    output_archive.writestr(info, payload)

            with self.assertRaisesRegex(docx_review_editor.DocxReviewError, "unknown.*w99"):
                docx_review_editor.patch_docx_reviews(
                    future,
                    root / "unsafe-edit.docx",
                    [{"op": "add_comment", "find": "target", "comment": "Review"}],
                )

    def test_tracked_replace_accept_and_reject_views(self) -> None:
        import docx

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = self._source(root, "Approve old wording")
            redline = root / "redline.docx"
            accepted = root / "accepted.docx"
            rejected = root / "rejected.docx"

            docx_review_editor.patch_docx_reviews(source, redline, [{
                "op": "tracked_replace",
                "find": "old",
                "replace": "new",
                "author": "Editor",
            }])
            with zipfile.ZipFile(redline) as archive:
                document_xml = archive.read("word/document.xml")
            self.assertIn(b"<w:del", document_xml)
            self.assertIn(b"<w:ins", document_xml)
            self.assertIn(b"delText", document_xml)

            docx_review_editor.patch_docx_reviews(redline, accepted, [{"op": "accept_changes"}])
            docx_review_editor.patch_docx_reviews(redline, rejected, [{"op": "reject_changes"}])
            self.assertEqual("Approve new wording", docx.Document(accepted).paragraphs[0].text)
            self.assertEqual("Approve old wording", docx.Document(rejected).paragraphs[0].text)
            with zipfile.ZipFile(accepted) as archive:
                self.assertNotIn(b"<w:ins", archive.read("word/document.xml"))
                self.assertNotIn(b"<w:del", archive.read("word/document.xml"))

    def test_bookmark_field_content_control_and_protection_are_native_word_objects(self) -> None:
        import docx

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "structured.docx"
            document = docx.Document()
            document.add_paragraph("Bookmark target")
            document.add_paragraph("Field placeholder")
            document.add_paragraph("Controlled text")
            document.save(source)
            output = root / "structured-output.docx"

            result = docx_review_editor.patch_docx_reviews(source, output, [
                {"op": "add_bookmark", "find": "Bookmark", "bookmarkName": "DecisionAnchor"},
                {"op": "insert_field", "find": "placeholder", "instruction": "REF DecisionAnchor", "displayText": "Bookmark"},
                {"op": "wrap_content_control", "find": "Controlled", "tag": "decision", "title": "Decision", "lock": "content"},
                {"op": "set_protection", "mode": "trackedChanges"},
            ])

            self.assertEqual(
                ["word/document.xml", "word/settings.xml"],
                result["changedParts"],
            )
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml")
                settings_xml = archive.read("word/settings.xml")
            self.assertIn(b"bookmarkStart", document_xml)
            self.assertIn(b"DecisionAnchor", document_xml)
            self.assertIn(b"fldSimple", document_xml)
            self.assertIn(b"REF DecisionAnchor", document_xml)
            self.assertIn(b"sdtContent", document_xml)
            self.assertIn(b"sdtContentLocked", document_xml)
            self.assertIn(b"documentProtection", settings_xml)
            self.assertIn(b"trackedChanges", settings_xml)

    def test_template_binding_targets_content_control_tags_and_bookmark_names(self) -> None:
        import docx

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "template.docx"
            structured = root / "structured.docx"
            bound = root / "bound.docx"
            document = docx.Document()
            document.add_paragraph("Customer placeholder")
            document.add_paragraph("Period placeholder")
            document.save(source)
            docx_review_editor.patch_docx_reviews(source, structured, [
                {
                    "op": "wrap_content_control",
                    "find": "Customer placeholder",
                    "tag": "customer_name",
                },
                {
                    "op": "add_bookmark",
                    "find": "Period placeholder",
                    "bookmarkName": "ReportPeriod",
                },
            ])
            result = docx_review_editor.patch_docx_reviews(structured, bound, [{
                "op": "bind_template",
                "bindings": {
                    "customer_name": "Nexa Labs",
                    "ReportPeriod": "2026 H2",
                },
            }])
            details = result["operations"][0]["detail"]["bindings"]
            self.assertEqual({"content_control", "bookmark"}, {item["target"] for item in details})
            with zipfile.ZipFile(bound) as archive:
                document_xml = archive.read("word/document.xml")
                self.assertIn(b"Nexa Labs", document_xml)
                self.assertIn(b"2026 H2", document_xml)
                self.assertNotIn(b"Customer placeholder", document_xml)
                self.assertNotIn(b"Period placeholder", document_xml)

    def test_field_allowlist_rejects_external_or_executable_instructions(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = self._source(root, "Unsafe field")
            with self.assertRaisesRegex(docx_review_editor.DocxReviewError, "safe.*allowlist"):
                docx_review_editor.patch_docx_reviews(source, root / "unsafe.docx", [{
                    "op": "insert_field",
                    "find": "Unsafe",
                    "instruction": "INCLUDETEXT https://example.invalid/secret",
                }])


if __name__ == "__main__":
    unittest.main()
